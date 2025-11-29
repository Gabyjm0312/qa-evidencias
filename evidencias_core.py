
import argparse
import os
from pathlib import Path
from datetime import datetime

import pandas as pd

# Dependencias opcionales, pero requeridas para generar DOCX
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception as e:
    raise SystemExit("Falta dependencia 'python-docx'. Instala con: pip install python-docx\nError: %s" % e)


DEFAULT_MAP = {
    "id": "Id caso de prueba",
    "ciclo": "Ciclo de Prueba",
    "sistema": "Sistema",
    "modulo": "Modulo a probar",
    "objetivo": "Objetivo/Descripción",
    "datos": "Datos de Prueba",
    "tipo": "Tipo de Prueba",
    "prereq": "Pre-requisitos",
    "pasos": "Pasos para Ejecución",
    "esperado": "Resultado Esperado",
    "obtenido": "Resultado Obtenido",
    "estado": "Estado de Prueba",
    "sla": "SLA",
    "fecha": "Fecha de Ejecución de Prueba",
    "tiempo": "Tiempo de Prueba (HH:MM)"
}


def _col(df, name):
    """Devuelve la serie para la columna 'name' con matching flexible (case-insensitive)."""
    for c in df.columns:
        if str(c).strip().lower() == str(name).strip().lower():
            return df[c].astype(str).fillna("")
    # Fallback: retorna serie vacía
    return pd.Series([""] * len(df))


def read_mapping(path_json: Path | None) -> dict:
    if path_json and path_json.exists():
        import json
        with open(path_json, "r", encoding="utf-8") as f:
            m = json.load(f)
        # mezcla con defaults
        merged = DEFAULT_MAP.copy()
        merged.update(m or {})
        return merged
    return DEFAULT_MAP.copy()


def steps_from_text(txt: str) -> list[str]:
    if not txt:
        return []
    # Normaliza saltos y divide por líneas no vacías
    parts = [s.strip() for s in str(txt).replace("\r", "").split("\n")]
    return [p for p in parts if p]


def add_kv_table(doc: Document, kv: list[tuple[str, str]]):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for k, v in kv:
        row = table.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v if v is not None else "")
    return table


def sanitize_filename(s: str, default="Caso"):
    s = s.strip() if s else default
    safe = []
    for ch in s[:80]:
        safe.append(ch if ch.isalnum() or ch in " .,_-()" else "_")
    out = "".join(safe).strip() or default
    return out


def build_single_doc(row: dict, out_file: Path):
    doc = Document()

    # Título
    h = doc.add_heading(f"Evidencia - Caso de Prueba {row.get('id') or ''}", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Objetivo/Descripción
    obj = row.get("objetivo") or ""
    if obj:
        doc.add_paragraph(obj)

    # Tabla de metadatos
    add_kv_table(doc, [
        ("Ciclo de Prueba", row.get("ciclo", "")),
        ("Sistema", row.get("sistema", "")),
        ("Módulo", row.get("modulo", "")),
        ("Tipo de Prueba", row.get("tipo", "")),
        ("Estado", row.get("estado", "")),
        ("SLA", row.get("sla", "")),
        ("Fecha de Ejecución", row.get("fecha", "")),
        ("Tiempo (HH:MM)", row.get("tiempo", "")),
        ("Datos de Prueba", row.get("datos", "")),
        ("Pre-requisitos", row.get("prereq", "")),
    ])

    # Pasos
    doc.add_heading("Pasos", level=1)
    pasos_list = steps_from_text(row.get("pasos", ""))
    if pasos_list:
        for s in pasos_list:
            doc.add_paragraph(s)
            
    else:
        doc.add_paragraph("(Sin pasos definidos en el Excel)")

    # Resultados
    doc.add_heading("Resultado Esperado", level=1)
    doc.add_paragraph(row.get("esperado") or "(Completar)")

    doc.add_heading("Resultado Obtenido", level=1)
    doc.add_paragraph(row.get("obtenido") or "(Completar)")

    # Evidencias (área para arrastrar imágenes)
    doc.add_heading("Evidencias", level=1)
    doc.add_paragraph("Inserta aquí las imágenes de evidencia (arrastrar y soltar en Word).")
    t2 = doc.add_table(rows=1, cols=1)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "⬇️ Arrastra y suelta tus imágenes aquí"

    doc.save(out_file)


def build_consolidated(rows: list[dict], out_file: Path):
    doc = Document()
    title = doc.add_heading("Evidencias de Casos de Prueba", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")

    for i, row in enumerate(rows):
        if i > 0:
            doc.add_page_break()
        doc.add_heading(f"Caso de Prueba {row.get('id') or ''}", level=1)
        obj = row.get("objetivo") or ""
        if obj:
            doc.add_paragraph(obj)

        add_kv_table(doc, [
            ("Ciclo de Prueba", row.get("ciclo", "")),
            ("Sistema", row.get("sistema", "")),
            ("Módulo", row.get("modulo", "")),
            ("Tipo de Prueba", row.get("tipo", "")),
            ("Estado", row.get("estado", "")),
            ("SLA", row.get("sla", "")),
            ("Fecha de Ejecución", row.get("fecha", "")),
            ("Tiempo (HH:MM)", row.get("tiempo", "")),
            ("Datos de Prueba", row.get("datos", "")),
            ("Pre-requisitos", row.get("prereq", "")),
        ])

        doc.add_paragraph("")
        doc.add_heading("Pasos para Ejecución", level=2)
        pasos_list = steps_from_text(row.get("pasos", ""))
        if pasos_list:
            for s in pasos_list:
                 doc.add_paragraph(s)
                
        else:
            doc.add_paragraph("(Sin pasos definidos en el Excel)")

        doc.add_heading("Resultado Esperado", level=2)
        doc.add_paragraph(row.get("esperado") or "(Completar)")
        doc.add_heading("Resultado Obtenido", level=2)
        doc.add_paragraph(row.get("obtenido") or "(Completar)")

        doc.add_heading("Evidencias", level=2)
        doc.add_paragraph("Inserta aquí las imágenes de evidencia (arrastrar y soltar en Word).")
        t2 = doc.add_table(rows=1, cols=1)
        t2.style = "Table Grid"
        t2.rows[0].cells[0].text = "⬇️ Arrastra y suelta tus imágenes aquí"

    doc.save(out_file)


def parse_args():
    ap = argparse.ArgumentParser(description="Generar documentos de evidencia desde Excel.")
    ap.add_argument("--excel", required=False, help="Ruta al Excel con los casos de prueba. Si se omite, se detecta automáticamente el primer .xls/.xlsx en la carpeta actual.")
    ap.add_argument("--hoja", default="Casos de Prueba", help="Nombre de la hoja. Default: 'Casos de Prueba'.")
    ap.add_argument("--salida", default="./evidencias_out", help="Carpeta de salida. Default: ./evidencias_out")
    ap.add_argument("--sin-consolidado", action="store_true", help="No generar DOCX consolidado.")
    ap.add_argument("--sin-individuales", action="store_true", help="No generar DOCX individuales por caso.")
    ap.add_argument("--map", dest="map_json", default=None, help="JSON con mapeo de columnas si difieren de los nombres por defecto.")
    return ap.parse_args()



import glob
def _autodetect_excel(cwd: Path, recursive=False):
    patterns = ["*.xlsx", "*.xls"]
    files = []
    if recursive:
        for p in patterns:
            files.extend(sorted(cwd.rglob(p), key=lambda p: p.stat().st_mtime, reverse=True))
    else:
        for p in patterns:
            files.extend(sorted(cwd.glob(p), key=lambda p: p.stat().st_mtime, reverse=True))
    return files[0] if files else None

def _pick_sheet_name(xls, requested: str | None):
    import re
    if requested:
        return requested
    # Heurística: busca algo parecido a "Casos de Prueba"
    for sh in xls.sheet_names:
        s = sh.strip().lower()
        if "casos" in s and "prueb" in s:
            return sh
    # Si no hay match, usa la primera
    return xls.sheet_names[0] if xls.sheet_names else None

def _run(a):
    """
    Lógica principal reutilizable. Antes estaba dentro de main().
    Recibe un objeto 'a' con los mismos atributos que los argumentos de línea de comandos.
    """

    # Autodetección del Excel si no se pasa --excel
    if a.excel:
        excel_path = Path(a.excel)
    else:
        excel_path = _autodetect_excel(Path.cwd(), recursive=False)
        if not excel_path:
            raise SystemExit("No se encontró ningún archivo .xls/.xlsx en la carpeta actual. Pasa --excel o coloca el archivo junto al script.")
        print(f"[INFO] Excel autodetectado: {excel_path}")

    if not excel_path.exists():
        raise SystemExit(f"No existe el Excel: {excel_path}")

    # Si no se indicó hoja, intenta detectar por heurística
        # Si no se indicó hoja, intenta detectar por heurística
    try:
        xls = pd.ExcelFile(excel_path)
    except Exception as e:
        raise SystemExit(f"No se pudo abrir el Excel. Error: {e}")

    # Determinar hoja
    hoja_name = _pick_sheet_name(xls, a.hoja if a.hoja else None)
    # 🔴 IMPORTANTE: cerramos el ExcelFile para liberar el archivo en Windows
    xls.close()

    if not hoja_name:
        raise SystemExit("No fue posible determinar la hoja a leer. Usa --hoja para especificarla.")

    try:
        df = pd.read_excel(excel_path, sheet_name=hoja_name, dtype=str).fillna("")
        print(f"[INFO] Hoja utilizada: {hoja_name}")
    except Exception as e:
        raise SystemExit(f"No se pudo leer la hoja '{hoja_name}'. Error: {e}")

    # mapeo de columnas
    mapping = read_mapping(Path(a.map_json) if a.map_json else None)

    rows = []
    for i in range(len(df)):
        row = {
            k: _col(df, mapping[k]).iloc[i].strip() if mapping.get(k) else ""
            for k in mapping.keys()
        }
        # Asegura un id inteligible
        if not row["id"]:
            row["id"] = f"{i+1:03d}"
        rows.append(row)

    # Si no se personaliza salida, por defecto crea junto al Excel
    out_dir = Path(a.salida)
    if a.salida == "./evidencias_out":
        out_dir = excel_path.parent / "evidencias_out"
    out_dir.mkdir(parents=True, exist_ok=True)

    # Genera consolidado
    if not a.sin_consolidado:
        consolidated = out_dir / "Evidencias_Consolidadas.docx"
        build_consolidated(rows, consolidated)
        print(f"[OK] Consolidado: {consolidated}")

    # Genera individuales
    if not a.sin_individuales:
        folder = out_dir / "Casos"
        folder.mkdir(parents=True, exist_ok=True)
        for r in rows:
            name_fragment = sanitize_filename(r.get("objetivo") or "Caso")
            path = folder / f"{r['id']} - {name_fragment}.docx"
            build_single_doc(r, path)
        print(f"[OK] Individuales en: {folder}")

    print("[LISTO] Proceso completado.")


def generar_evidencias_desde_excel(
    excel_path: str,
    hoja: str = "Casos de Prueba",
    salida: str = "./evidencias_out",
    sin_consolidado: bool = False,
    sin_individuales: bool = False,
    map_json: str | None = None,
):
    """
    Envoltura para usar este script desde otro código (API, app web, etc.).
    En vez de leer parámetros desde consola, recibimos argumentos directos.
    """
    class Args:
        pass

    a = Args()
    a.excel = excel_path
    a.hoja = hoja
    a.salida = salida
    a.sin_consolidado = sin_consolidado
    a.sin_individuales = sin_individuales
    a.map_json = map_json

    _run(a)


def main():
    """Entry point de línea de comandos."""
    a = parse_args()
    _run(a)


if __name__ == "__main__":
    main()
